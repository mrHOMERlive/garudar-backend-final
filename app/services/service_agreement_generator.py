"""Генератор Service Agreement DOCX из локального шаблона.

Алгоритм — зеркало `nda_generator.py` (preserve formatting через `char_to_run`
map). Используется `app/template/service_agreement/SA_ENG_V1.docx`.

Соответствие плейсхолдеров шаблона полям модели `ServiceAgreementRequest`:

| Шаблон                | Поле модели           |
|-----------------------|-----------------------|
| [EFFECTIVE_DATE]      | effective_date        |
| [COMPANY_NAME]        | company_name          |
| [COUNTRY]             | country               |
| [ADDRESS]             | address               |
| [SIGNATORY_NAME]      | signatory_name        |
| [SIGNATORY_TITLE]     | signatory_title       |
| [REGISTRATION_NUMBER] | registration_number   |
| [TAX_ID]              | tax_id                |
| [CONTACT_EMAIL]       | contact_email         |
| [CONTACT_PHONE]       | contact_phone         |
| [TERM]                | term                  |

Если какого-то плейсхолдера нет в .docx — значение просто не подставляется
(it's not an error). При добавлении новых плейсхолдеров шаблон редактируется
вручную в Word.
"""

from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Dict

from docx import Document


TEMPLATE_NAME = "SA_ENG_V1.docx"
TEMPLATE_PATH = Path(__file__).parent.parent / "template" / "service_agreement" / TEMPLATE_NAME

FIELD_MAPPING: Dict[str, str] = {
    "[EFFECTIVE_DATE]":      "effective_date",
    "[COMPANY_NAME]":        "company_name",
    "[COUNTRY]":             "country",
    "[ADDRESS]":             "address",
    "[SIGNATORY_NAME]":      "signatory_name",
    "[SIGNATORY_TITLE]":     "signatory_title",
    "[REGISTRATION_NUMBER]": "registration_number",
    "[TAX_ID]":              "tax_id",
    "[CONTACT_EMAIL]":       "contact_email",
    "[CONTACT_PHONE]":       "contact_phone",
    "[TERM]":                "term",
}


def _format_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return str(value)


def _replace_placeholders(doc: Document, fields: Dict) -> None:
    replacements = {
        placeholder: _format_value(fields[field_name])
        for placeholder, field_name in FIELD_MAPPING.items()
        if field_name in fields and fields[field_name] is not None
    }

    if not replacements:
        return

    def replace_in_paragraph(paragraph) -> None:
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        modified_text = full_text
        for placeholder, value in replacements.items():
            modified_text = modified_text.replace(placeholder, value)

        if modified_text == full_text:
            return

        char_to_run = []
        for run_idx, run in enumerate(paragraph.runs):
            char_to_run.extend([run_idx] * len(run.text))

        position_map = []
        old_pos = 0

        sorted_placeholders = sorted(replacements.keys(), key=len, reverse=True)

        while old_pos < len(full_text):
            placeholder_found = None
            for placeholder in sorted_placeholders:
                if full_text[old_pos: old_pos + len(placeholder)] == placeholder:
                    placeholder_found = placeholder
                    break

            if placeholder_found:
                replacement = replacements[placeholder_found]
                run_idx = char_to_run[old_pos] if old_pos < len(char_to_run) else 0
                for _ in range(len(replacement)):
                    position_map.append(run_idx)
                old_pos += len(placeholder_found)
            else:
                run_idx = char_to_run[old_pos] if old_pos < len(char_to_run) else 0
                position_map.append(run_idx)
                old_pos += 1

        new_run_texts = [""] * len(paragraph.runs)
        for i, char in enumerate(modified_text):
            if i < len(position_map):
                new_run_texts[position_map[i]] += char
            else:
                new_run_texts[0] += char

        for i, run in enumerate(paragraph.runs):
            run.text = new_run_texts[i]

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def generate(fields: Dict) -> bytes:
    """Сгенерировать заполненный SA DOCX. Возвращает bytes."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Шаблон Service Agreement не найден: {TEMPLATE_PATH}"
        )

    doc = Document(str(TEMPLATE_PATH))
    _replace_placeholders(doc, fields)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
