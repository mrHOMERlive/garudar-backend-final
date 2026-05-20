"""
Генератор NDA (Non-Disclosure Agreement) DOCX из локального шаблона.

Подход и алгоритм скопированы из `service_agreement_generator.py`:
- Шаблон лежит локально в `app/template/nda/NDA_ENG_V1.docx`.
- Реквизиты «Disclosing Party» (PT GAN) hardcoded в .docx — их не
  подставляем, чтобы текст договора оставался юридически выверенным.
- Подставляются только данные «Receiving Party» (партнёра) — те поля,
  которые клиент заполняет в форме `ClientNDA.jsx`.

Текущий шаблон унаследован от mitra (PT MITRA - NDA_eng.docx) и
содержит плейсхолдеры формата `[POINT N]`. Перед production-релизом
заменить hardcoded реквизиты PT MITRA на PT GAN в самом .docx через
Word (адрес, регистрационный номер, имя подписанта).

Соответствие плейсхолдеров шаблона полям модели `NdaRequest`:

| Шаблон    | Поле модели                |
|-----------|----------------------------|
| [POINT 1] | effective_date             |
| [POINT 2] | partner_name_en            |
| [POINT 3] | partner_country_en         |
| [POINT 4] | partner_inn                |
| [POINT 5] | partner_signatory_en       |
| [POINT 5.1] | partner_signatory_title_en |
| [POINT 6] | partner_address_en         |
| [POINT 7] | partner_contact_email      |
"""

from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Dict

from docx import Document


TEMPLATE_NAME = "NDA_ENG_V1.docx"
TEMPLATE_PATH = Path(__file__).parent.parent / "template" / "nda" / TEMPLATE_NAME

# Маппинг плейсхолдеров шаблона -> ключи полей из словаря fields.
# Длинные плейсхолдеры (`[POINT 5.1]`) сортируются раньше коротких
# в `_replace_placeholders` (см. `sorted(..., key=len, reverse=True)`),
# чтобы `[POINT 5]` не съел префикс `[POINT 5.1]`.
FIELD_MAPPING: Dict[str, str] = {
    "[POINT 1]":   "effective_date",
    "[POINT 2]":   "partner_name_en",
    "[POINT 3]":   "partner_country_en",
    "[POINT 4]":   "partner_inn",
    "[POINT 5.1]": "partner_signatory_title_en",
    "[POINT 5]":   "partner_signatory_en",
    "[POINT 6]":   "partner_address_en",
    "[POINT 7]":   "partner_contact_email",
}


def _format_value(value) -> str:
    """Приводит значение к строке для подстановки в DOCX."""
    if value is None:
        return ""
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return str(value)


def _replace_placeholders(doc: Document, fields: Dict) -> None:
    """
    Заменяет плейсхолдеры в документе на значения из fields.
    Обрабатывает параграфы и ячейки таблиц, сохраняя форматирование runs.

    Алгоритм copy-paste из `service_agreement_generator.py:_replace_placeholders`:
    собираем `char_to_run` map, строим `position_map` для каждого
    символа результирующей строки и распределяем символы обратно по runs
    с сохранением исходного форматирования.
    """
    replacements = {
        placeholder: _format_value(fields[field_name])
        for placeholder, field_name in FIELD_MAPPING.items()
        if field_name in fields and fields[field_name] is not None
    }

    if not replacements:
        return

    def replace_in_paragraph(paragraph) -> None:
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        modified_text = full_text
        for placeholder, value in replacements.items():
            modified_text = modified_text.replace(placeholder, value)

        if modified_text == full_text:
            return

        char_to_run = []
        for run_idx, run in enumerate(paragraph.runs):
            char_to_run.extend([run_idx] * len(run.text))

        position_map = []
        old_pos = 0

        sorted_placeholders = sorted(replacements.keys(), key=len, reverse=True)

        while old_pos < len(full_text):
            placeholder_found = None
            for placeholder in sorted_placeholders:
                if full_text[old_pos: old_pos + len(placeholder)] == placeholder:
                    placeholder_found = placeholder
                    break

            if placeholder_found:
                replacement = replacements[placeholder_found]
                run_idx = char_to_run[old_pos] if old_pos < len(char_to_run) else 0
                for _ in range(len(replacement)):
                    position_map.append(run_idx)
                old_pos += len(placeholder_found)
            else:
                run_idx = char_to_run[old_pos] if old_pos < len(char_to_run) else 0
                position_map.append(run_idx)
                old_pos += 1

        new_run_texts = [""] * len(paragraph.runs)
        for i, char in enumerate(modified_text):
            if i < len(position_map):
                new_run_texts[position_map[i]] += char
            else:
                new_run_texts[0] += char

        for i, run in enumerate(paragraph.runs):
            run.text = new_run_texts[i]

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def generate(fields: Dict) -> bytes:
    """
    Генерирует заполненный NDA из локального шаблона.

    Args:
        fields: словарь с данными для подстановки. Ключи соответствуют
            значениям FIELD_MAPPING (partner_name_en, effective_date, ...).
            Лишние ключи игнорируются; отсутствующие — оставляют
            плейсхолдер незаменённым (что заметно при визуальной проверке
            DOCX и позволяет легко обнаружить пропущенное поле).

    Returns:
        bytes — содержимое .docx файла, готовое для отдачи клиенту или
        загрузки в S3.

    Raises:
        FileNotFoundError: если шаблон не найден на диске.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Шаблон NDA не найден: {TEMPLATE_PATH}"
        )

    doc = Document(str(TEMPLATE_PATH))

    _replace_placeholders(doc, fields)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
